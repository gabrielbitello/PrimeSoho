import os
import uuid
from num2words import num2words
from docx import Document
import re

counter_dict = {}
subcounter_dict = {}

def criar_tabela_ao_redor_do_valor(paragrafo, valor):
    """Cria uma tabela simples ao redor do valor, adicionando o valor à célula da tabela."""
    # Cria uma tabela com uma linha e uma coluna (ao redor do valor)
    tabela = paragrafo.add_paragraph().add_run().add_table(rows=1, cols=1)

    # Adiciona o valor à célula da tabela
    celula = tabela.cell(0, 0)
    celula.text = valor

    # Aplica estilo à tabela (bordas internas apenas)
    tabela.style = 'Table Grid'

    # Retorna o parágrafo com a tabela ao redor do valor
    return paragrafo

def converter_numero_para_texto(valor):
    """Converte um número para texto em português."""
    if valor is None:
        return ""
    try:
        return num2words(int(valor), lang='pt_BR')
    except ValueError:
        return ""

def substituir_texto(paragrafo, campo, valor):
    """Substitui todas as ocorrências de {campo} no parágrafo por valor."""
    campo_formatado = f"{{{campo}}}"
    if campo_formatado in paragrafo.text:
        novo_texto = paragrafo.text.replace(campo_formatado, str(valor) if valor else "")
        paragrafo.clear()
        paragrafo.add_run(novo_texto)

def parse_string(input_str):
    if ':' in input_str:
        x, y = input_str.split(':')
        return x, y
    else:
        x = input_str
        y = None
        return x, y

def aplicar_regras_para_valor(valor, yaml_data, campo_verificado, doc, dados):
    """Aplica as regras de forma automática com base nas configurações no YAML para o valor da variável."""
    
    # Se não houver regras, retornamos o valor sem alteração
    if not yaml_data.get('regras'):
        return valor

    # Obtém as regras para o campo verificado, se houver
    regras_aplicadas = yaml_data.get('regras', [])

    # Iterando sobre as regras (agora assumimos que é uma lista)
    for regra_obj in regras_aplicadas:

        if isinstance(regra_obj, dict):
            for regra, regra_valor in regra_obj.items():

                # Exemplo de regra para adicionar uma caixa (não implementado, mas podemos colocar o valor dentro de uma tabela ou similar)
                if regra == "Add_box" and regra_valor:
                    # Aqui você pode adicionar a lógica para adicionar o valor à tabela ou o que for necessário
                    inc = "inc"  # Exemplo, mas pode ser qualquer ação necessária

                # Converte número para texto, se a regra for válida
                elif regra == "Number_To_Text" and isinstance(regra_valor, str):
                    valor_para_converter = dados.get(regra_valor, None)
                    if valor_para_converter is not None:
                        valor = converter_numero_para_texto(valor_para_converter)

                # Aplica contagem a partir de um valor usando regex (exemplo de contador)
                elif regra == "Counter" and isinstance(regra_valor, str):
                    XY_value = regra_valor
                    if XY_value:
                        x, y = parse_string(XY_value)
                        counter_value = get_counter_value(x, y)
                        valor = counter_value + valor  # Ajustado para somar o contador ao valor
    return valor

def get_counter_value(x, y=None):
    x = int(x) 
    y = int(y) if y is not None else None  
    # Se X ainda não foi inicializado, ele começa em 1 (sem incrementar ainda)
    if x not in counter_dict:
        counter_dict[x] = 1
        subcounter_dict[x] = {}

    # Se Y está presente, não incrementa X, mas usa o valor atual dele e gerencia Y
    if y is not None:
        if y not in subcounter_dict[x]:
            subcounter_dict[x][y] = 1  
        else:
            subcounter_dict[x][y] += 1 
        return f"{counter_dict[x] - 1}.{subcounter_dict[x][y]}"  

    # Se Y não está presente, apenas incrementa X e retorna o valor atualizado
    valor_atual = counter_dict[x]  
    counter_dict[x] += 1  
    return str(valor_atual)

def substituir_variaveis_no_paragrafo(paragrafo, dados, yaml_data, doc):
    """Substitui variáveis no parágrafo conforme as regras, garantindo que cada contador seja atualizado corretamente."""

    texto_original = paragrafo.text

    # Expressão regular para encontrar padrões como {counter:x} ou {counter:x:y}
    pattern = r"\{counter:(\d+)(?::(\d+))?\}"

    # Função auxiliar para substituição dinâmica
    def substituir_match(match):
        x = match.group(1)
        y = match.group(2)
        

        # Obtém o novo valor do contador
        novo_valor = get_counter_value(x, y) if y else get_counter_value(x)

        return str(novo_valor)

    # Substituir diretamente nos 'runs' sem processar todo o parágrafo antes
    for run in paragrafo.runs:
        if "{counter:" in run.text:  # Só processa os runs que realmente contêm um contador
            run.text = re.sub(pattern, substituir_match, run.text)

    # Processa outras variáveis além de counter
    for chave, valor in dados.items():
        chave_formatada = f"{{{chave}}}"  # Variáveis são referenciadas como {variavel}

        if chave_formatada in texto_original:

            # Obtém a configuração específica do YAML
            documentos_config = yaml_data.get('Documentos', {}).get('Documentos-Config', [])
            item_filtrado = next((item for item in documentos_config if item.get('nome') == chave), None)

            # Verifica se há condições para essa variável e se elas são atendidas
            if not verificar_condicoes(dados, item_filtrado, chave):
                # Se a condição falhar, remove a variável do parágrafo
                for run in paragrafo.runs:
                    if chave_formatada in run.text:
                        run.text = run.text.replace(chave_formatada, '')  # Remover a variável
                continue  # Passa para a próxima variável

            # Verifica se o valor da variável ainda está vazio
            novo_valor = valor or ''
            if not novo_valor:
                # Caso o valor esteja vazio, atribuir o primeiro valor do campo "variáveis" no YAML
                variaveis_yaml = item_filtrado.get('variaveis', {})
                if variaveis_yaml:
                    novo_valor = variaveis_yaml[0]  # Atribui o primeiro valor de 'variaveis'

            # Aplica regras antes de retornar o valor
            novo_valor = aplicar_regras_para_valor(novo_valor, item_filtrado, chave, doc, dados)

            # Mantém a formatação e substitui apenas a variável nos 'runs'
            for run in paragrafo.runs:
                if chave_formatada in run.text:
                    run.text = run.text.replace(chave_formatada, str(novo_valor))

    return paragrafo




def substituir_variaveis_nas_tabelas(doc, dados, yaml_data):
    """Substitui variáveis dentro de tabelas do documento."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragrafo in cell.paragraphs: 
                    substituir_variaveis_no_paragrafo(paragrafo, dados, yaml_data, doc)

def verificar_condicoes(dados, yaml_data, campo_verificado):
    """Verifica as condições de um campo, buscando no YAML as condições associadas ao campo."""

    # Obtém as condições do YAML relacionadas ao campo
    condicao = yaml_data.get('condicao', {})

    if not condicao:
        return True  # Se não houver condição, o campo é considerado válido.

    if not isinstance(condicao, dict):
        return False

    # Itera sobre as condições para o campo
    for chave, valor in condicao.items():
        
        if chave in dados:
            campo_valor = dados[chave]

            # Verifica se o valor da condição é uma lista
            if isinstance(valor, list):
                # Se o campo está em uma lista de valores aceitos, passa a condição
                if campo_valor not in valor:
                    return False
            # Verifica a condição booleana
            elif isinstance(valor, bool):
                if valor:  # Espera que o campo tenha um valor
                    if not campo_valor:
                        return False
                else:  # Espera que o campo NÃO tenha valor
                    if campo_valor:
                        return False
            # Caso a condição seja um valor simples (não booleano nem lista)
            elif campo_valor != valor:  # Verifica se o valor do campo é o esperado
                return False
        else:
            return False

    return True  # Se todas as condições foram atendidas

def gen_docx(dados, folder, yaml_data):
    """Preenche o modelo DOCX com os dados do formulário e do YAML."""
    caminho_template = os.path.abspath(os.path.join('app', 'docs', folder, f'{folder}.docx'))

    unique_id = str(uuid.uuid4())[:8]

    # Adiciona variáveis do YAML aos dados, se ainda não existirem
    for campo_yaml in yaml_data.get('Documentos', {}).get('Documentos-Config', []):
        nome_campo = campo_yaml.get('nome')
        if nome_campo and nome_campo not in dados:
            dados[nome_campo] = campo_yaml.get('variaveis', [None])[0] or ""

    # Carrega o arquivo DOCX
    doc = Document(caminho_template)

    if not doc.paragraphs:
        raise ValueError(f"O template {caminho_template} não contém parágrafos.")

    # Substituição de variáveis no documento
    for paragrafo in doc.paragraphs:
        substituir_variaveis_no_paragrafo(paragrafo, dados, yaml_data, doc)

    # Substituição dentro das tabelas
    substituir_variaveis_nas_tabelas(doc, dados, yaml_data)

    # Salva o arquivo gerado
    caminho_saida = os.path.abspath(f'./app/output/{folder}_preenchido_{unique_id}.docx')
    doc.save(caminho_saida)
    return caminho_saida
