import os
import uuid
from num2words import num2words
from docx import Document
import re
from copy import deepcopy
import logging
from django.core.exceptions import ValidationError
from typing import Optional, Union, Dict, Any, List



# Constants
PARENT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_PATH = os.path.join(PARENT_DIR, 'docs')
OUTPUT_PATH = os.path.join(PARENT_DIR, 'output')

# Global Dictionaries (Consider using classes or a more encapsulated approach if this becomes more complex)
COUNTER_DICT: Dict[int, int] = {}
SUBCOUNTER_DICT: Dict[int, Dict[int, int]] = {}

COUNTERS: Dict[str, int] = {}  # Initialize the counters dictionary


def clear_counters():
    global COUNTER_DICT, SUBCOUNTER_DICT, COUNTERS
    COUNTER_DICT.clear()
    SUBCOUNTER_DICT.clear()
    COUNTERS.clear()

# Get the logger
logger = logging.getLogger('juridico')

def create_table_around_value(paragraph, value: str):
    """
    Creates a simple table around the given value within a Word document paragraph.

    Args:
        paragraph: The docx.paragraph.Paragraph object where the table will be inserted.
        value: The string value to be placed inside the table.

    Returns:
        The original paragraph, for chaining or further manipulation.
    """
    table = paragraph.add_paragraph().add_run().add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = value
    table.style = 'Table Grid'
    return paragraph


def convert_number_to_text(value: Optional[Union[int, float, str]]) -> str:
    """
    Converts a numerical value to its textual representation in Brazilian Portuguese.

    Args:
        value: The numerical value to convert. Can be an integer, float, or string.  None is accepted and returns "".

    Returns:
        The textual representation of the number, or an empty string if conversion fails.
    """
    if value is None:
        return ""
    try:
        value_str = str(value).replace('.', '').replace(' ', '')
        return num2words(int(value_str), lang='pt_BR')
    except ValueError:
        return ""

def replace_text(paragraph, field: str, value: str):
    """
    Replaces all occurrences of {field} in the paragraph with the given value.

    Args:
        paragraph: The docx.paragraph.Paragraph object to modify.
        field: The string representing the placeholder to replace (e.g., 'name').
        value: The string value to insert in place of the placeholder.
    """
    formatted_field = f"{{{field}}}"
    if formatted_field in paragraph.text:
        new_text = paragraph.text.replace(formatted_field, str(value) if value else "")
        paragraph.clear()
        paragraph.add_run(new_text)

def parse_string(input_str: str) -> tuple[str, Optional[str], Optional[str]]:
    """
    Parses a string to extract components based on delimiters ":" and "/".

    Args:
        input_str: The string to parse (e.g., "x:y/modifier" or "x/modifier" or "x").

    Returns:
        A tuple containing x, y (optional), and modifier (optional).
    """
    if ":" in input_str:
        x, rest = input_str.split(":", 1)
        if "/" in rest:
            y, modifier = rest.split("/")
        else:
            y = rest
            modifier = None
        return x, y, modifier
    elif "/" in input_str:
        x, modifier = input_str.split("/")
        return x, None, modifier
    else:
        return input_str, None, None

def apply_rules_to_value(value: Any, yaml_data: Dict[str, Any], field_verified: str, doc: Document, data: Dict[str, Any], yaml: Dict[str, Any], index: str | None = None, row: object | None = None, table: object | None = None) -> Any:
    """
    Applies rules to a value based on configurations in the YAML file.

    Args:
        value: The initial value to be processed.
        yaml_data: The YAML data containing rules and configurations.
        field_verified: The field being verified (used for context in rules).
        doc: The Document object (used if rules require document context).
        data: The data dictionary containing values for formatting.
        yaml: The entire YAML configuration for broader context access.
        index (str | None): O índice, que pode ser uma string ou None.
        row (object | None): The table row object (if the value is in a table).
        table (object | None): The table object (if the value is in a table).
    Returns:
        The modified value after applying all relevant rules.
    """
    if not yaml_data.get('regras'):
        return value

    for rule_obj in yaml_data.get('regras', []):
        if isinstance(rule_obj, dict):
            for rule, rule_value in rule_obj.items():
                if rule == "Add_box" and rule_value:
                    pass  # Placeholder for future action
                elif rule == "Number_To_Text" and isinstance(rule_value, str):
                    if rule_value.startswith('#'):
                        rule_value = f"form-{index}-{rule_value[1:]}" if index is not None else f"form-{rule_value[1:]}"
                    value_to_convert = data.get(rule_value)
                    
                    # Se o valor for None, aplicar regras adicionais
                    if value_to_convert is None:
                        # Buscar regras para o campo
                        document_configs = yaml.get('Documentos', {}).get('Documentos-Config', [])
                        field_config = next((item for item in document_configs if item.get('nome') == rule_value.replace('form-', '').split('-')[-1]), None)
                        
                        if field_config and field_config.get('regras'):
                            # Aplicar regras ao valor nulo
                            value_to_convert = apply_rules_to_value(None, field_config, rule_value, doc, data, yaml, index)
                    
                    if value_to_convert is not None:
                        value = convert_number_to_text(value_to_convert)
                elif rule == "Formater" and isinstance(rule_value, str):
                    value = format_text_with_data(rule_value, data, yaml, field_verified, index)
                elif rule == "Counter" and isinstance(rule_value, str):
                    value = handle_counter_rule(rule_value, value)
                elif rule == "Clear" and rule_value and value is not None:
                    value = Clear_box(row, table, value)
                elif rule == "Sum" and isinstance(rule_value, list) and len(rule_value) >= 2:
                    value = total_sum(data, rule_value, yaml, doc, index)

    return value

def total_sum(data: Dict[str, Any], rule_value: List[str], yaml: Dict[str, Any], doc: Document, index: str | None = None ) -> float:
    # Nova regra: Sum - soma dois ou mais valores
                    total_sum = 0
                    for field_to_sum in rule_value:
                        # Ajustar o nome do campo se necessário
                        if field_to_sum.startswith('#'):
                            field_to_sum = f"form-{index}-{field_to_sum[1:]}" if index is not None else f"form-{field_to_sum[1:]}"
                        
                        # Obter o valor do campo
                        field_value = data.get(field_to_sum)
                        
                        # Se o valor for None ou não numérico, tentar aplicar regras
                        if field_value is None or not isinstance(field_value, (int, float)):
                            # Buscar regras para o campo
                            document_configs = yaml.get('Documentos', {}).get('Documentos-Config', [])
                            field_name = field_to_sum.replace('form-', '').split('-')[-1]
                            field_config = next((item for item in document_configs if item.get('nome') == field_name), None)
                            
                            if field_config and field_config.get('regras'):
                                # Aplicar regras ao valor
                                processed_value = apply_rules_to_value(field_value, field_config, field_to_sum, doc, data, yaml, index)
                                try:
                                    # Tentar converter para número
                                    if processed_value is not None:
                                        if isinstance(processed_value, str):
                                            # Remover caracteres não numéricos (exceto ponto decimal)
                                            processed_value = re.sub(r'[^\d.]', '', processed_value)
                                        field_value = float(processed_value)
                                except (ValueError, TypeError):
                                    field_value = 0
                        
                        # Adicionar à soma total
                        try:
                            if field_value is not None:
                                total_sum += float(field_value)
                        except (ValueError, TypeError):
                            # Se não puder converter para número, ignorar
                            pass
                    
                    # Atualizar o valor com a soma total
                    return str(total_sum) if total_sum else ""

def Clear_box(row: object | None, table: object | None, value: Any) -> Any:
    """
    Limpa a célula da tabela se o valor for falso.

    Args:
        row (object | None): The table row object (if the value is in a table).
        table (object | None): The table object (if the value is in a table).
        value: The value to check for clearing.

    Returns:
        The value if it is not empty, or an empty string if it is.
    """
    if row is not None and table is not None:
        row._element.clear_content()
        remove_row_below(table, row)
    return value

def format_text_with_data(format_string: str, data: Dict[str, Any], yaml: Dict[str, Any], field: str, index: str | None = None) -> str:
    """
    Formats a string using values from the provided data dictionary and YAML configurations.

    Args:
        format_string: The string containing placeholders to be replaced.
        data: The dictionary containing data values.
        yaml: The YAML configuration for accessing document-level configurations.
        field: The field name (não utilizado nesta implementação, mas mantido para compatibilidade).
        index (str | None): O índice, que pode ser uma string ou None.

    Returns:
        The formatted string.
    """
    def replace_variable(match):
        key = match.group(1)
        oldkey = key

        if key.startswith('#'):
            oldkey = key[1:]
            key = f"form-{index}-{key[1:]}" if index is not None else f"form-{key[1:]}"

        if key in data and data[key] is not None:
            return str(data[key])
        else:
            document_configs = yaml.get('Documentos', {}).get('Documentos-Config', [])
            item_filtered = next((item for item in document_configs if item.get('nome') == oldkey), None)

            if item_filtered:
                if item_filtered.get('regras') or item_filtered.get('condicao'):
                    if not verify_conditions(data, item_filtered, None, index):
                        return match.group(0)

                    variables_yaml = item_filtered.get('variaveis', [])
                    new_value = variables_yaml[0] if variables_yaml else None

                    if new_value is None:
                        new_value = apply_rules_to_value(None, item_filtered, key, None, data, yaml, index)

                    if new_value is not None:
                        return str(new_value)

        # Se chegou aqui, não foi possível substituir o valor
        print(f"Não foi possível substituir a tag: {match.group(0)}")
        return match.group(0)

    return re.sub(r'{(.*?)}', replace_variable, format_string)

def handle_counter_rule(rule_value: str, value: Any) -> str:
    """
    Handles the counter rule by parsing the rule value and updating the counter.

    Args:
        rule_value: The string containing the counter rule (e.g., "x:y/B").
        value: The current value to which the counter value will be prepended.

    Returns:
        The updated value with the counter prepended.
    """
    if rule_value:
        x, y, modifier = parse_string(rule_value)
        counter_value = get_counter_value(x, y)
        if modifier == "B":
            counter_value = f"**{counter_value} - **"
        return counter_value + str(value)  # Ensure value is converted to string

    return str(value)  # Ensure value is converted to string

def get_counter_value(x: str, y: Optional[str] = None) -> str:
    """
    Retrieves and increments counter values based on x and y identifiers.

    Args:
        x: The primary counter identifier.
        y: The secondary counter identifier (optional).

    Returns:
        The string representation of the counter value.
    """
    x_int = int(x)
    y_int = int(y) if y is not None else None

    if x_int not in COUNTER_DICT:
        COUNTER_DICT[x_int] = 1
        SUBCOUNTER_DICT[x_int] = {}

    if y_int is not None:
        if y_int not in SUBCOUNTER_DICT[x_int]:
            SUBCOUNTER_DICT[x_int][y_int] = 1
        else:
            SUBCOUNTER_DICT[x_int][y_int] += 1
        return f"{COUNTER_DICT[x_int] - 1}.{SUBCOUNTER_DICT[x_int][y_int]}"

    if y == 'Y' and x == '0':
        x_int = 1

    current_value = COUNTER_DICT[x_int]
    COUNTER_DICT[x_int] += 1

    return str(current_value)

def replace_match(match: re.Match) -> str:
    """
    Replaces a regex match with a counter value.

    Args:
        match: The regex match object.

    Returns:
        The string representation of the new counter value.
    """
    x = match.group(1)
    y = match.group(2)
    new_value = get_counter_value(x, y) if y else get_counter_value(x)
    return str(new_value)

def process_data(data: Dict[str, Any], parsed_data_options: Dict[str, Any]) -> Dict[str, Any]:
    """
    Processes data based on parsed data options to create new data entries.

    Args:
        data: The original data dictionary.
        parsed_data_options: Configuration options for processing data.

    Returns:
        A new dictionary with processed data.
    """
    new_data = {}
    counters = {}

    print("Processing data")
    for name, config in parsed_data_options.items():
        if name in data:
            value = data[name]
            if isinstance(value, list):
                for i, item in enumerate(value):
                    if isinstance(item, dict):
                        for key, val in item.items():
                            if key not in counters:
                                counters[key] = 0
                            else:
                                counters[key] += 1
                            new_key = f"form-{counters[key]}-{key}"
                            new_data[new_key] = val
                    else:
                        new_key = f"form-{i}-{name}"
                        new_data[new_key] = item
            elif isinstance(value, dict):
                for key, val in value.items():
                    if key not in counters:
                        counters[key] = 0
                    else:
                        counters[key] += 1
                    new_key = f"form-{counters[key]}-{key}"
                    new_data[new_key] = val
            else:
                new_key = f"form-0-{name}"
                new_data[new_key] = value

    return new_data

def process_special_inputs(original_text: str) -> str:
    """
    Processes special inputs in the text to format them correctly, incrementing counters for repeated names.

    Args:
        original_text: The original text containing special inputs (e.g., "{#Nome}").

    Returns:
        The processed text with formatted special inputs (e.g., "{form-0-Nome}", "{form-1-Nome}").
    """

    def replace_input(match: re.Match) -> str:
        """
        Replaces a match with the formatted input, incrementing the counter if the name is repeated.
        """
        name: str = match.group(1)  # Extract the name from the match
        if name.startswith('form-'):
            # If already in the format {form-X-Nome}, do not alter it
            return match.group(0)

        if name not in COUNTERS:
            # If the name is encountered for the first time, initialize its counter
            COUNTERS[name] = 0
        else:
            # If the name has appeared before, increment its counter
            COUNTERS[name] += 1

        # Return the new formatted string
        return f"{{form-{COUNTERS[name]}-{name}}}"

    # Pattern to capture {#Nome} or {any-prefix-Nome} but prioritize {#Nome}
    pattern = r'{(?:#|\w+-)(\w+)}'

    processed_text: str = re.sub(pattern, replace_input, original_text)

    return processed_text

def extract_name(key: str) -> str:
    """
    Extracts the name from a formatted key string.

    Args:
        key: The formatted key string (e.g., "form-0-name").

    Returns:
        The extracted name (e.g., "name").
    """
    parts = key.split('-')
    return '-'.join(parts[2:])

def replace_counter(match: re.Match) -> str:
    x = match.group(1)
    y = match.group(2)
    return get_counter_value(x, y) if y else get_counter_value(x)

# Compile the regex pattern for counters
counter_pattern = re.compile(r"{counter:(\d+)(?::(\d+))?}")

def replace_variables_in_paragraph(paragraph: object, data: Dict[str, Any], yaml_data: Dict[str, Any], doc: object, parsed_data_options: Dict[str, Any], row: object = None, table_o: object = None) -> object:
    original_text = paragraph.text
    processed_text = process_special_inputs(original_text)
    
    # Handle counters first
    processed_text = counter_pattern.sub(replace_counter, processed_text)
    
    # Pré-processamento dos dados e configurações
    replacements = {}
    yaml_cache = {}
    document_configs = yaml_data.get('Documentos', {}).get('Documentos-Config', [])
    
    # Encontrar todas as chaves no texto
    all_keys = re.findall(r'{([^}]+)}', processed_text)
    
    for key in all_keys:
        formatted_key = f"{{{key}}}"
        
        # Busca flexível da chave
        matching_key = next((k for k in data.keys() if key in k), None)
        
        if matching_key:
            value = data[matching_key]
            reformatted_key = extract_name(matching_key) if matching_key.startswith("form") else matching_key
            index = matching_key.split('-')[1] if matching_key.startswith("form") else None
            
            if reformatted_key not in yaml_cache:
                yaml_cache[reformatted_key] = next((item for item in document_configs if item.get('nome') == reformatted_key), None)
            
            item_filtered = yaml_cache[reformatted_key]
            
            if item_filtered and not verify_conditions(data, item_filtered, reformatted_key, index):
                replacements[formatted_key] = ''
                continue
            
            new_value = str(value) if value is not None else ''
            if not new_value and item_filtered and item_filtered.get('variaveis'):
                new_value = str(item_filtered['variaveis'][0]) if item_filtered['variaveis'][0] is not None else ''
            
            if item_filtered:
                new_value = apply_rules_to_value(new_value, item_filtered, reformatted_key, doc, data, yaml_data, index)
            
            replacements[formatted_key] = str(new_value)
            
            if not new_value:
                logger.warning(f"Campo '{key}' encontrado (correspondendo a '{matching_key}'), mas com valor vazio.")
        else:
            logger.error(f"Campo '{key}' não encontrado nos dados. Chaves similares: {[k for k in data.keys() if key in k]}")
    
    # Substituição em lote
    for old, new in replacements.items():
        processed_text = processed_text.replace(old, new)
    
    # Aplicar o texto processado ao parágrafo
    paragraph.clear()
    if "**" in processed_text:
        parts = processed_text.split("**")
        for i, part in enumerate(parts):
            run = paragraph.add_run(part)
            if i % 2 == 1:
                run.bold = True
    else:
        paragraph.add_run(processed_text)
    
    # Lidar com tabelas vazias
    if table_o is not None and not processed_text.strip():
        for row in table_o.rows:
            if all(not cell.text.strip() for cell in row.cells):
                table_o._tbl.remove(row._tr)
                break
    
    return paragraph

def replace_variables_in_tables(doc: Document, data: Dict[str, Any], yaml_data: Dict[str, Any], parsed_data_options: Dict[str, Any]):
    """Replaces variables inside tables of the document."""
    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    replace_variables_in_paragraph(paragraph, data, yaml_data, doc, parsed_data_options, row, table)

def replace_variables_in_footers(doc: Document, data: Dict[str, Any], yaml_data: Dict[str, Any], parsed_data_options: Dict[str, Any]):
    """Replaces variables in the footers of the document."""
    for section in doc.sections:
        for footer in section.footer.paragraphs:
            replace_variables_in_paragraph(footer, data, yaml_data, doc, parsed_data_options)

def replace_variables_in_headers(doc: Document, data: Dict[str, Any], yaml_data: Dict[str, Any], parsed_data_options: Dict[str, Any]):
    """Replaces variables in the headers of the document."""
    for section in doc.sections:
        for header in section.header.paragraphs:
            replace_variables_in_paragraph(header, data, yaml_data, doc, parsed_data_options)


def verify_conditions(data: Dict[str, Any], yaml_data: Dict[str, Any], field_verified: str, index: str | None = None) -> bool:
    """
    Verifies conditions for a field based on YAML configurations, with advanced formset verification.

    Args:
        data: The dictionary containing data values.
        yaml_data: The YAML configuration data.
        field_verified: The field being verified.
        index: The index for formset fields, if applicable.

    Returns:
        True if all conditions are met or no conditions are specified, False otherwise.
    """
    try:
        condition = yaml_data.get('condicao', {})

        if not condition:
            return True

        if not isinstance(condition, dict):
            return False

        for key, value in condition.items():
            # Processar condições especiais para formsets
            if key.startswith('formset:'):
                parts = key.split(':')
                
                # Formato básico: formset:campo
                if len(parts) == 2:
                    field_name = parts[1]
                    # Encontrar todos os campos do formset
                    formset_values = [data[k] for k in data.keys() 
                                        if re.match(f"form-\d+-{field_name}$", k) and k in data]
                    
                    if isinstance(value, list):
                        # Verificar se pelo menos um campo tem um dos valores da lista
                        if not any(fv in value for fv in formset_values):
                            return False
                    else:
                        # Verificar se pelo menos um campo tem o valor específico
                        if value not in formset_values:
                            return False
                
                # Formato avançado: formset:campo:operador:valor
                elif len(parts) >= 3:
                    field_name = parts[1]
                    operator = parts[2]
                    
                    # Encontrar todos os campos do formset
                    formset_values = [data[k] for k in data.keys() 
                                        if re.match(f"form-\d+-{field_name}$", k) and k in data]
                    
                    if operator == "any":
                        # Verificar se qualquer campo tem o valor especificado
                        if isinstance(value, list):
                            if not any(fv in value for fv in formset_values):
                                return False
                        else:
                            if value not in formset_values:
                                return False
                    
                    elif operator == "all":
                        # Verificar se todos os campos têm o valor especificado
                        if isinstance(value, list):
                            if not all(fv in value for fv in formset_values):
                                return False
                        else:
                            if not all(fv == value for fv in formset_values):
                                return False
                    
                    elif operator == "count":
                        # Verificar se existe um número específico de campos com o valor
                        if isinstance(value, dict):
                            # Formato: {"A": 2, "B": 3} - precisa de 2 campos com valor A e 3 com valor B
                            for val, count in value.items():
                                actual_count = sum(1 for fv in formset_values if fv == val)
                                if actual_count < count:
                                    return False
                        else:
                            # Formato simples: precisa de pelo menos N campos com qualquer valor não vazio
                            actual_count = sum(1 for fv in formset_values if fv)
                            if actual_count < int(value):
                                return False
                    
                    elif operator == "combination":
                        # Verificar combinações específicas de valores
                        # Exemplo: formset:campo:combination:{"A": 1, "B": 1}
                        # Significa: precisa de pelo menos 1 campo com valor A E pelo menos 1 campo com valor B
                        if isinstance(value, dict):
                            for val, min_count in value.items():
                                actual_count = sum(1 for fv in formset_values if fv == val)
                                if actual_count < min_count:
                                    return False
                
                continue  # Prosseguir para a próxima condição após verificar o formset
            
            # Processamento regular para campos não-formset (código existente)
            adjusted_key = key
            
            # Check if index is not None and adjust key accordingly
            if index is not None:
                adjusted_key = f"form-{index}-{key}"

            if '/' in adjusted_key:
                keys = adjusted_key.split('/')
                condition_met = False

                for key_part in keys:
                    if key_part in data:
                        field_value = data[key_part]

                        if isinstance(value, list):
                            if field_value in value:
                                condition_met = True
                                break
                        elif isinstance(value, bool):
                            if value and not field_value:
                                return False
                            elif not value and field_value:
                                return False
                        else:
                            if isinstance(field_value, str) and isinstance(value, str):
                                if field_value == value:
                                    condition_met = True
                                    break
                            elif field_value == value:
                                condition_met = True
                                break

                if not condition_met:
                    return False

            else:
                if adjusted_key in data:
                    field_value = data[adjusted_key]

                    if isinstance(value, list):
                        if field_value not in value:
                            return False
                    elif isinstance(value, bool):
                        if value and not field_value:
                            return False
                        elif not value and field_value:
                            return False
                    else:
                        if isinstance(field_value, str) and isinstance(value, str):
                            if field_value != value:
                                return False
                        elif field_value != value:
                            return False

                else:
                    return False

        return True
    except Exception as e:
        print(f"Error in verify_conditions: {e}")
        raise

def insert_table_after(original_table: object, new_table: object):
    """Inserts a new table after the original table in the document."""
    original_table_element = original_table._element
    new_table_element = new_table._element
    original_table_element.addnext(new_table_element)

def insert_row_below(table: object, original_row: object):
    """Inserts a new row below the original row in the table."""
    new_row = table.add_row()

    for i, cell in enumerate(original_row.cells):
        new_row.cells[i].text = ""

def remove_row_below(table: object, original_row: object):
    """Remove a linha abaixo da linha original na tabela."""
    row_index = table.rows.index(original_row)
    if row_index + 1 < len(table.rows):
        table._tbl.remove(table.rows[row_index + 1]._tr)

def process_data_single(data: Dict[str, Any], parsed_data_options: Dict[str, Any], processed_fields: Dict[str, bool]) -> Dict[str, Any]:
    """
    Processa um único campo de dados com base nas opções analisadas, mantendo um histórico dos campos processados.

    Args:
        data: O dicionário de dados original.
        parsed_data_options: Opções de configuração para processamento de dados.
        processed_fields: Dicionário que mantém o registro dos campos já processados.

    Returns:
        Um novo dicionário com o campo processado e o histórico atualizado.
    """
    new_data = {}
    counters = {}

    print("Processando dados")
    
    # Encontra o primeiro campo não processado
    field_to_process = next((name for name in parsed_data_options if name in data and not processed_fields.get(name, False)), None)
    
    if field_to_process:
        value = data[field_to_process]
        if isinstance(value, list):
            for i, item in enumerate(value):
                if isinstance(item, dict):
                    for key, val in item.items():
                        if key not in counters:
                            counters[key] = 0
                        else:
                            counters[key] += 1
                        new_key = f"form-{counters[key]}-{key}"
                        new_data[new_key] = val
                else:
                    new_key = f"form-{i}-{field_to_process}"
                    new_data[new_key] = item
        elif isinstance(value, dict):
            for key, val in value.items():
                if key not in counters:
                    counters[key] = 0
                else:
                    counters[key] += 1
                new_key = f"form-{counters[key]}-{key}"
                new_data[new_key] = val
        else:
            new_key = f"form-0-{field_to_process}"
            new_data[new_key] = value
        
        # Marca o campo como processado
        processed_fields[field_to_process] = True
    
    return new_data, processed_fields


def generate_docx(data: Dict[str, Any], folder: str, yaml_data: Dict[str, Any], parsed_data_options: Dict[str, Any]) -> str:
    """
    Generates a DOCX file by populating a template with data from a form and YAML configurations.

    Args:
        data: The dictionary containing data to populate the document.
        folder: The folder containing the DOCX template.
        yaml_data: The YAML data containing document configurations.
        parsed_data_options: Options for parsing and processing data.

    Returns:
        The path to the generated DOCX file.
    """
    try:
        logger.info(f"Iniciando geração do documento para a pasta: {folder}")
        logger.info(f"Dados iniciais: {data}")
        logger.info(f"Opções de dados analisadas: {parsed_data_options}")

        template_path = os.path.abspath(os.path.join(DOCS_PATH, folder, f'{folder}.docx'))
        unique_id = str(uuid.uuid4())[:8]
        logger.info(f"Caminho do template: {template_path}")

        # Add YAML variables to data if they don't already exist
        for yaml_field in yaml_data.get('Documentos', {}).get('Documentos-Config', []):
            field_name = yaml_field.get('nome')
            if field_name and field_name not in data:
                if not yaml_field.get('grupo'):
                    data[field_name] = None
                    logger.info(f"Adicionado campo YAML: {field_name}")

        # Load the DOCX file
        try:
            doc = Document(template_path)
            logger.info("Documento carregado com sucesso")
        except FileNotFoundError:
            logger.error(f"Template file not found: {template_path}")
            raise FileNotFoundError(f"Template file not found: {template_path}")
        except Exception as e:
            logger.error(f"Error loading document: {e}")
            raise Exception(f"Error loading document: {e}")

        # Data processing
        logger.info("Iniciando processamento de dados")
        processed_data = process_data(data, parsed_data_options)
        logger.info(f"Dados processados: {processed_data}")
        data.update(processed_data)
        logger.info(f"Dados atualizados após processamento: {data}")

        # Table cloning logic
        logger.info("Iniciando lógica de clonagem de tabelas")
        for group in parsed_data_options:
            if group in data:
                search_text = parsed_data_options[group].get('buscador')
                multiplier = data[group]
                logger.info(f"Processando grupo: {group}, buscador: {search_text}, multiplicador: {multiplier}")
                
                num_dicts = sum(1 for item in multiplier if isinstance(item, dict))
                logger.info(f"Número de dicionários: {num_dicts}")

                for yaml_field in yaml_data.get('Documentos', {}).get('Documentos-Config', []):
                    field_name = yaml_field.get('nome')
                    if field_name and yaml_field.get('grupo'):
                        for i in range(num_dicts):
                            new_field_name = f"form-{i}-{field_name}"
                            if new_field_name not in data:
                                data[new_field_name] = None
                                logger.info(f"Campo adicionado: {new_field_name}")

                if num_dicts > 1:
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if search_text in cell.text:
                                    original_table = table
                                    logger.info(f"Tabela encontrada com texto de busca: {search_text}")
                                    for _ in range(num_dicts - 1):
                                        new_table = deepcopy(original_table)
                                        insert_table_after(original_table, new_table)
                                        logger.info("Nova tabela inserida")
                                    break
            else:
                logger.info(f"Grupo não encontrado nos dados: {group}")

        # Remove grupos originais após o processamento
        for group in parsed_data_options:
            if group in data:
                del data[group]
                logger.info(f"Grupo removido: {group}")

        if not doc.paragraphs:
            logger.error(f"The template {template_path} contains no paragraphs.")
            raise ValueError(f"The template {template_path} contains no paragraphs.")

        # Variable substitution
        for paragraph in doc.paragraphs:
            replace_variables_in_paragraph(paragraph, data, yaml_data, doc, parsed_data_options)

        replace_variables_in_tables(doc, data, yaml_data, parsed_data_options)
        replace_variables_in_footers(doc, data, yaml_data, parsed_data_options)
        replace_variables_in_headers(doc, data, yaml_data, parsed_data_options)

        


        # Save the generated file
        output_filename = f'{folder}_preenchido_{unique_id}.docx'
        output_path_formatted = os.path.abspath(os.path.join(OUTPUT_PATH, output_filename))

        # Certifique-se de que o diretório existe
        if not os.path.exists(OUTPUT_PATH):
            os.makedirs(OUTPUT_PATH, exist_ok=True)
            logger.info(f"Diretório 'output' criado em: {OUTPUT_PATH}")

        doc.save(output_path_formatted)
        
        clear_counters()

        return output_filename

    except Exception as e:
        clear_counters()
        logger.error(f"Error generating document: {str(e)}", exc_info=True)
        raise ValidationError(f"Error generating document: {str(e)}")